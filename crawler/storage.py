"""
数据储存模块 - 支持 JSON/CSV/SQLite/Markdown/TXT/DOCX 格式
"""

import json
import csv
import re
import logging
import sqlite3
from pathlib import Path
from typing import List, Dict, Any, Optional


class Storage:
    """通用数据存储器"""

    def __init__(self, config: dict = None):
        self.config = config or {}
        self.default_format = self.config.get("storage", {}).get("default", "auto")
        self.logger = logging.getLogger("storage")
        # 预编译图片URL正则
        self._img_pattern = re.compile(
            r'https?://[^\s]+\.(jpg|jpeg|png|gif|webp|bmp|svg)',
            re.IGNORECASE
        )
        # 预编译HTML标签正则（语义识别）
        self._html_tag_pattern = re.compile(r'<[^>]+>', re.IGNORECASE)
        # 预编译Markdown语法正则（语义识别）
        self._md_syntax_pattern = re.compile(
            r'^#{1,6}\s|^\*\s|^\d+\.\s|^\>\s|\*\*|__|\[.+?\]\(.+?\)|```', re.MULTILINE
        )
        # auto模式配置（config.yaml → storage.auto）
        auto_cfg = self.config.get("storage", {}).get("auto", {})
        self.auto_prefer = auto_cfg.get("prefer", "balanced")  # balanced / readability / structured / weighted
        self.disabled_formats = set(auto_cfg.get("disabled_formats", []))
        self.long_text_threshold = auto_cfg.get("long_text_threshold", 500)
        self.max_readable_records = auto_cfg.get("max_readable_records", 200)
        self.sample_size = auto_cfg.get("sample_size", 30)
        self.csv_fill_rate_threshold = auto_cfg.get("csv_fill_rate_threshold", 0.8)
        self.csv_special_char_threshold = auto_cfg.get("csv_special_char_threshold", 0.05)
        # 显式初始化，避免 _save_docx/_generate_file_index_md 中 getattr 兜底
        self._download_info = []

    # 文件名后缀 → 格式映射
    _EXT_MAP = {
        ".json": "json", ".csv": "csv", ".md": "markdown", ".markdown": "markdown",
        ".txt": "txt", ".docx": "docx", ".db": "sqlite", ".sqlite": "sqlite", ".sqlite3": "sqlite",
    }
    # 格式降级链
    _DEGRADATION = {
        "docx": ["docx", "markdown", "txt", "json"],
        "markdown": ["markdown", "txt", "json"],
        "csv": ["csv", "json"],
        "txt": ["txt", "json"],
        "json": ["json"],
        "sqlite": ["sqlite", "json"],
    }

    def save(self, data: List[Dict[str, Any]], output_path: str, fmt: str = None,
             dual_output: bool = False) -> str:
        """保存数据到文件（支持文件名后缀优先、格式降级兜底、混合输出）"""
        output = Path(output_path)

        # 如果路径是已存在的目录，在其中生成文件名
        if output.is_dir():
            from datetime import datetime
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            output = output / f"爬取结果_{ts}"
            self.logger.info(f"输出路径是目录，自动生成文件名: {output.name}")

        output.parent.mkdir(parents=True, exist_ok=True)

        # 文件名后缀优先：带后缀的文件名直接用对应格式，跳过auto判断
        if fmt is None or fmt == "auto":
            ext_fmt = self._EXT_MAP.get(output.suffix.lower())
            if ext_fmt and ext_fmt not in self.disabled_formats:
                fmt = ext_fmt
                self.logger.debug(f"文件名后缀优先：{output.suffix} → {ext_fmt}")

        fmt = fmt or self.default_format

        # auto 模式 — 根据数据特征自动选择格式
        if fmt == "auto":
            fmt = self._auto_detect_format(data)

        # 过滤禁用格式
        if fmt in self.disabled_formats:
            self.logger.debug(f"格式 {fmt} 被禁用，切换到 auto 重新检测")
            fmt = self._auto_detect_format(data)

        # 如果文件名没有后缀，根据格式自动补全
        _fmt_ext = {"json": ".json", "csv": ".csv", "sqlite": ".db",
                    "markdown": ".md", "txt": ".txt", "docx": ".docx"}
        if not output.suffix and fmt in _fmt_ext:
            output = output.with_suffix(_fmt_ext[fmt])
            self.logger.debug(f"自动补全后缀: {output.name}")

        # 混合输出：同时保存 JSON（结构化备份）+ 可读格式
        if dual_output and fmt not in ("json", "sqlite"):
            json_path = output.with_suffix(".json")
            self._save_by_format(data, json_path, "json")
            self.logger.info(f"混合输出：JSON 备份已保存到 {json_path}")

        # 格式降级兜底链
        chain = self._DEGRADATION.get(fmt, [fmt, "json"])
        last_error = None
        for try_fmt in chain:
            if try_fmt in self.disabled_formats:
                continue
            try:
                return self._save_by_format(data, output, try_fmt)
            except Exception as e:
                last_error = e
                self.logger.warning(f"格式 {try_fmt} 保存失败: {e}，尝试降级...")
        raise RuntimeError(f"所有格式保存均失败，最后错误: {last_error}")

    def _save_by_format(self, data: List[Dict], path: Path, fmt: str) -> str:
        """按指定格式保存（内部方法）"""
        if fmt == "json":
            return self._save_json(data, path)
        elif fmt == "csv":
            return self._save_csv(data, path)
        elif fmt == "sqlite":
            return self._save_sqlite(data, path)
        elif fmt == "markdown":
            return self._save_markdown(data, path)
        elif fmt == "txt":
            return self._save_txt(data, path)
        elif fmt == "docx":
            return self._save_docx(data, path)
        else:
            raise ValueError(f"不支持的格式: {fmt}")

    def _auto_detect_format(self, data: List[Dict[str, Any]]) -> str:
        """决策树版：五层优先级递进，自动识别数据最优保存格式
        支持 config 配置驱动阈值、语义识别、加权评分模式、禁用格式过滤
        """
        # ========== 阶段一：输入校验 ==========
        if not data:
            self.logger.debug("自动格式判定：数据集为空，默认使用 json")
            return "json"

        record_count = len(data)
        first_keys = data[0].keys()

        # ========== 阶段二：均匀采样 ==========
        sz = min(self.sample_size, record_count)
        if record_count <= sz:
            sample = data
        else:
            step = record_count // sz
            sample = [data[i * step] for i in range(sz)]

        # ========== 阶段三：多维度特征统计 ==========
        has_nested = False
        total_text_len = 0
        long_text_count = 0
        image_count = 0
        all_fields_set = set()
        total_field_times = 0
        non_empty_count = 0
        special_char_count = 0
        html_tag_count = 0
        md_syntax_count = 0

        for item in sample:
            all_fields_set.update(item.keys())
            for value in item.values():
                total_field_times += 1
                if isinstance(value, (list, dict)):
                    has_nested = True
                    continue
                if isinstance(value, str):
                    if value.strip():
                        non_empty_count += 1
                    text_len = len(value)
                    total_text_len += text_len
                    if text_len > self.long_text_threshold:
                        long_text_count += 1
                    if self._img_pattern.search(value):
                        image_count += 1
                    special_char_count += value.count(',') + value.count('\n')
                    # 语义识别：HTML标签占比
                    html_tag_count += len(self._html_tag_pattern.findall(value))
                    # 语义识别：Markdown语法检测
                    md_syntax_count += len(self._md_syntax_pattern.findall(value))

        # 衍生指标
        avg_text_len = total_text_len / max(1, len(sample))
        field_count = len(all_fields_set)
        fill_rate = non_empty_count / max(1, total_field_times)
        special_char_ratio = special_char_count / max(1, total_text_len)
        html_tag_ratio = html_tag_count / max(1, total_text_len)
        md_syntax_ratio = md_syntax_count / max(1, len(sample))

        # ========== 加权评分模式（可选）==========
        if self.auto_prefer == "weighted":
            return self._weighted_score_format(
                has_nested, avg_text_len, long_text_count, image_count,
                field_count, fill_rate, special_char_ratio, record_count,
                html_tag_ratio, md_syntax_ratio, first_keys
            )

        # ========== 阶段四：五层决策树（命中即返回）==========
        # 过滤禁用格式的辅助函数
        def pick(primary, fallback="json"):
            if primary not in self.disabled_formats:
                return primary
            if fallback not in self.disabled_formats:
                return fallback
            return "json"

        # 第1层：嵌套结构检测
        # 但如果是文章类数据（有title+text长文本），嵌套只是元数据，优先可读格式
        if has_nested:
            has_article_text = any(
                isinstance(item.get("text"), str) and len(item["text"]) > 200
                for item in sample
            )
            if has_article_text and record_count <= 100:
                result = pick("markdown", "json")
                self.logger.debug(f"自动格式判定：{result} | 依据：文章类数据（有长文本），嵌套为元数据")
                return result
            self.logger.debug(f"自动格式判定：json | 依据：嵌套结构且非文章类，记录数{record_count}")
            return "json"

        # 第2层：带图小批量 → DOCX（含降级 markdown）
        if image_count > 0 and record_count <= 50:
            try:
                import docx  # noqa: F401
                result = pick("docx", "markdown")
                self.logger.debug(f"自动格式判定：{result} | 依据：图片{image_count}个，记录数{record_count}")
                return result
            except ImportError:
                result = pick("markdown", "json")
                self.logger.debug(f"自动格式判定：{result} | 依据：有图片但无python-docx，降级")
                return result

        # 第3层：长文本可读性 → 小批量Markdown / 大批量JSON
        if avg_text_len > 300 or long_text_count >= 2:
            # 语义识别：内容含大量HTML标签 → 优先JSON保留原始结构
            if html_tag_ratio > 0.1:
                self.logger.debug(f"自动格式判定：json | 依据：HTML标签占比{html_tag_ratio:.1%}高，保留结构")
                return "json"
            # 语义识别：内容已含Markdown语法 → 优先Markdown
            if md_syntax_ratio > 0.3 and "markdown" not in self.disabled_formats:
                self.logger.debug(f"自动格式判定：markdown | 依据：检测到Markdown语法，记录数{record_count}")
                return "markdown"
            # prefer 模式影响
            if self.auto_prefer == "readability":
                if record_count <= self.max_readable_records and "markdown" not in self.disabled_formats:
                    return "markdown"
                return "json"
            # 默认balanced逻辑
            if record_count <= 100 and field_count <= 10:
                result = pick("markdown", "json")
                self.logger.debug(f"自动格式判定：{result} | 依据：长文本小批量，平均{avg_text_len:.0f}字符")
                return result
            else:
                self.logger.debug(f"自动格式判定：json | 依据：长文本但记录{record_count}>100")
                return "json"

        # 第4层：扁平表格场景 → CSV（需同时满足4项条件）
        fields_consistent = field_count == len(first_keys)
        all_short = avg_text_len <= 200
        fill_ok = fill_rate >= self.csv_fill_rate_threshold
        char_ok = special_char_ratio < self.csv_special_char_threshold

        if fields_consistent and all_short and fill_ok and char_ok:
            result = pick("csv", "json")
            self.logger.debug(
                f"自动格式判定：{result} | 依据：字段一致，填充率{fill_rate:.0%}，"
                f"平均{avg_text_len:.0f}，特殊字符{special_char_ratio:.1%}"
            )
            return result

        # 第5层：兜底 → JSON
        self.logger.debug(
            f"自动格式判定：json | 依据：未匹配专项规则（兜底），"
            f"记录{record_count}，字段{field_count}"
        )
        return "json"

    def _weighted_score_format(self, has_nested, avg_text_len, long_text_count, image_count,
                                field_count, fill_rate, special_char_ratio, record_count,
                                html_tag_ratio, md_syntax_ratio, first_keys) -> str:
        """加权评分制：多格式打分，取最高分（可扩展、可调权重）"""
        scores = {"json": 0, "csv": 0, "markdown": 0, "txt": 0, "docx": 0}

        # 一票否决：嵌套结构 → JSON直接胜出
        if has_nested:
            self.logger.debug("加权评分：嵌套结构 → json 一票否决")
            return "json"

        # 评分规则
        # 平均文本长度
        if avg_text_len > 300:
            scores["markdown"] += 3
            scores["txt"] += 2
            scores["csv"] -= 2
        elif avg_text_len < 50:
            scores["csv"] += 2
            scores["json"] += 1

        # 含图片URL且记录少
        if image_count > 0 and record_count <= 50:
            scores["docx"] += 5
            scores["markdown"] += 2

        # 字段整齐且填充率高
        fields_consistent = field_count == len(first_keys)
        if fields_consistent and fill_rate > 0.8:
            scores["csv"] += 3
            scores["json"] += 1

        # 记录数大
        if record_count > 200:
            scores["csv"] += 2
            scores["json"] += 2
            scores["markdown"] -= 2

        # 字段数多
        if field_count > 15:
            scores["json"] += 2
            scores["csv"] -= 1

        # 特殊字符多 → CSV减分
        if special_char_ratio > 0.05:
            scores["csv"] -= 3
            scores["markdown"] += 1

        # 语义识别：HTML标签多 → JSON加分
        if html_tag_ratio > 0.1:
            scores["json"] += 3
            scores["markdown"] -= 1

        # 语义识别：Markdown语法多 → Markdown加分
        if md_syntax_ratio > 0.3:
            scores["markdown"] += 3

        # 长文本字段多
        if long_text_count >= 2:
            scores["markdown"] += 2
            scores["txt"] += 1

        # 过滤禁用格式
        for f in self.disabled_formats:
            scores.pop(f, None)

        # 取最高分
        best = max(scores, key=scores.get)
        self.logger.debug(
            f"加权评分：{best} | 得分: {scores} | "
            f"avg_text={avg_text_len:.0f}, records={record_count}, fields={field_count}"
        )
        return best

    def load(self, input_path: str) -> List[Dict[str, Any]]:
        """从文件加载数据"""
        path = Path(input_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {input_path}")

        suffix = path.suffix.lower()
        if suffix == ".json":
            return self._load_json(path)
        elif suffix == ".csv":
            return self._load_csv(path)
        elif suffix in (".db", ".sqlite", ".sqlite3"):
            return self._load_sqlite(path)
        elif suffix in (".md", ".markdown"):
            return self._load_markdown(path)
        elif suffix == ".txt":
            return self._load_txt(path)
        else:
            # 默认尝试 JSON
            return self._load_json(path)

    def set_download_info(self, info: List[Dict[str, str]]):
        """设置下载文件信息（与下载模块联动，在输出中生成本地文件索引）
        info: [{"filename": "...", "url": "...", "local_path": "..."}, ...]
        """
        self._download_info = info

    def _generate_file_index_md(self) -> str:
        """生成本地文件索引的Markdown文本"""
        if not getattr(self, '_download_info', None):
            return ""
        lines = ["\n\n---\n\n## 下载文件索引\n"]
        lines.append("| 序号 | 文件名 | 原始URL | 本地路径 |")
        lines.append("|------|--------|---------|----------|")
        for i, f in enumerate(self._download_info, 1):
            fname = f.get("filename", "")
            url = f.get("url", "")
            local = f.get("local_path", "")
            lines.append(f"| {i} | {fname} | {url} | {local} |")
        return "\n".join(lines)

    def _save_json(self, data: List[Dict], path: Path) -> str:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        return str(path)

    def _save_csv(self, data: List[Dict], path: Path) -> str:
        if not data:
            path.write_text("", encoding="utf-8")
            return str(path)

        keys = list(data[0].keys())
        for item in data[1:]:
            for k in item.keys():
                if k not in keys:
                    keys.append(k)

        with open(path, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.DictWriter(f, fieldnames=keys, extrasaction="ignore")
            writer.writeheader()
            writer.writerows(data)
        return str(path)

    def _save_sqlite(self, data: List[Dict], path: Path) -> str:
        if not data:
            return str(path)

        conn = sqlite3.connect(str(path))
        cursor = conn.cursor()

        # 收集所有字段
        all_keys = []
        for item in data:
            for k in item.keys():
                if k not in all_keys:
                    all_keys.append(k)

        # 创建表
        columns = ", ".join(f'"{k}" TEXT' for k in all_keys)
        table_name = "data"
        cursor.execute(f"DROP TABLE IF EXISTS {table_name}")
        cursor.execute(f"CREATE TABLE {table_name} ({columns})")

        # 插入数据
        placeholders = ", ".join("?" * len(all_keys))
        for item in data:
            values = []
            for k in all_keys:
                v = item.get(k)
                if isinstance(v, (dict, list)):
                    v = json.dumps(v, ensure_ascii=False)
                values.append(str(v) if v is not None else None)
            cursor.execute(f"INSERT INTO {table_name} VALUES ({placeholders})", values)

        conn.commit()
        conn.close()
        return str(path)

    def _save_markdown(self, data: List[Dict], path: Path) -> str:
        if not data:
            path.write_text("", encoding="utf-8")
            return str(path)

        # 判断是否为文章类数据（有title+text长文本）
        first = data[0]
        is_article = (
            isinstance(first.get("text"), str) and len(first["text"]) > 200
            and isinstance(first.get("title"), str) and first["title"].strip()
        )

        if is_article and len(data) <= 10:
            # 文章类：输出可读的 Markdown 文档
            lines = []
            for i, item in enumerate(data):
                if i > 0:
                    lines.append("\n\n---\n")
                # 标题
                title = item.get("title", "无标题").strip()
                lines.append(f"# {title}\n")
                # URL
                url = item.get("url", "")
                if url:
                    lines.append(f"> 原文链接：{url}\n")
                # 正文
                text = item.get("text", "").strip()
                if text:
                    lines.append(text)
                # 元数据摘要
                meta = item.get("metadata", {})
                if isinstance(meta, dict) and meta:
                    desc = meta.get("description", "")
                    if desc:
                        lines.append(f"\n\n> **摘要**：{desc}")
                # 下载文件索引
                file_index = self._generate_file_index_md()
                if file_index:
                    lines.append(file_index)
            content = "\n".join(lines)
            path.write_text(content, encoding="utf-8")
            return str(path)

        # 表格类数据：输出 Markdown 表格
        lines = []
        keys = list(data[0].keys())
        lines.append("| " + " | ".join(keys) + " |")
        lines.append("| " + " | ".join(["---"] * len(keys)) + " |")

        # 数据行
        for item in data:
            row = []
            for k in keys:
                v = item.get(k, "")
                if isinstance(v, (dict, list)):
                    v = json.dumps(v, ensure_ascii=False)
                v = str(v).replace("|", "\\|").replace("\n", " ")
                row.append(v)
            lines.append("| " + " | ".join(row) + " |")

        # 下载文件索引（与下载模块联动）
        file_index = self._generate_file_index_md()
        if file_index:
            lines.append(file_index)

        path.write_text("\n".join(lines), encoding="utf-8")
        return str(path)

    def _save_txt(self, data: List[Dict], path: Path) -> str:
        """保存为纯文本格式 - 每条记录用分隔线隔开，字段名: 值"""
        if not data:
            path.write_text("", encoding="utf-8")
            return str(path)

        lines = []
        separator = "=" * 60
        
        for i, item in enumerate(data):
            if i > 0:
                lines.append("")
                lines.append(separator)
                lines.append("")
            
            for key, value in item.items():
                if isinstance(value, (dict, list)):
                    value = json.dumps(value, ensure_ascii=False, indent=2)
                elif value is None:
                    value = "(空)"
                else:
                    value = str(value)
                lines.append(f"{key}: {value}")

        path.write_text("\n".join(lines), encoding="utf-8")
        return str(path)

    def _save_docx(self, data: List[Dict], path: Path) -> str:
        """保存为Word文档格式（支持嵌入图片）"""
        import os
        import tempfile
        try:
            import requests as _requests
        except ImportError:
            _requests = None
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise ImportError(
                "需要安装 python-docx 库才能导出Word文档。\n"
                "请运行: pip install python-docx"
            )

        if not data:
            doc = Document()
            doc.save(str(path))
            return str(path)

        doc = Document()

        # 不加总标题，直接输出记录内容
        doc.add_paragraph(f'共 {len(data)} 条记录')
        doc.add_paragraph('')

        # 临时目录存放下载的图片
        tmp_dir = tempfile.mkdtemp(prefix="crawler_img_")
        img_count = 0

        try:
            for i, item in enumerate(data):
                doc.add_heading(f'记录 {i + 1}', level=1)

                # ── 标题 ──
                item_title = item.get("title", "")
                if item_title:
                    item_title = str(item_title).replace("\n", " ").replace("\r", "").replace("\t", " ")
                    doc.add_heading(item_title, level=2)

                # ── 来源/日期等元信息 ──
                meta_line_parts = []
                for mk in ("url", "author", "date"):
                    v = item.get(mk)
                    if v:
                        v = str(v).replace("\n", " ").replace("\r", "").replace("\t", " ")
                        meta_line_parts.append(f"{mk}: {v}")
                if meta_line_parts:
                    mp = doc.add_paragraph(" | ".join(meta_line_parts))
                    mp.style.font.size = Pt(9)

                # ── 正文 ──
                text = item.get("text")
                if text:
                    # 按段落分割写入
                    for para_text in str(text).split("\n\n"):
                        para_text = para_text.strip()
                        if para_text:
                            doc.add_paragraph(para_text)

                # ── 图片（下载并嵌入） ──
                images = item.get("images")
                if images and isinstance(images, list):
                    doc.add_paragraph('')
                    doc.add_heading('图片', level=3)
                    for idx, img_url in enumerate(images):
                        if not isinstance(img_url, str) or not img_url.startswith("http"):
                            continue
                        img_saved = False
                        if _requests:
                            try:
                                resp = _requests.get(img_url, timeout=15, headers={
                                    "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"
                                })
                                if resp.status_code == 200 and len(resp.content) > 500:
                                    # 判断扩展名
                                    ct = resp.headers.get("Content-Type", "").lower()
                                    if "png" in ct or img_url.lower().endswith(".png"):
                                        ext = ".png"
                                    elif "gif" in ct or img_url.lower().endswith(".gif"):
                                        ext = ".gif"
                                    elif "webp" in ct or img_url.lower().endswith(".webp"):
                                        ext = ".webp"
                                    else:
                                        ext = ".jpg"
                                    img_path = os.path.join(tmp_dir, f"img_{img_count}{ext}")
                                    with open(img_path, "wb") as f:
                                        f.write(resp.content)
                                    # 嵌入 Word（限制宽度 6 英寸）
                                    doc.add_picture(img_path, width=Inches(6))
                                    img_count += 1
                                    img_saved = True
                            except Exception:
                                pass
                        if not img_saved:
                            # 下载失败则写 URL
                            doc.add_paragraph(f"[图片下载失败] {img_url}")

                # ── 其他字段（metadata, links 等） ──
                skip_keys = {"title", "text", "images", "html"}
                for key, value in item.items():
                    if key in skip_keys:
                        continue
                    if value is None or value == [] or value == {}:
                        continue
                    p = doc.add_paragraph()
                    run = p.add_run(f'{key}: ')
                    run.bold = True
                    if isinstance(value, (dict, list)):
                        value = json.dumps(value, ensure_ascii=False, indent=2)
                    else:
                        value = str(value)
                    p.add_run(value)

                # 分隔线
                if i < len(data) - 1:
                    doc.add_paragraph('_' * 50)

        finally:
            # 清理临时图片
            import shutil
            try:
                shutil.rmtree(tmp_dir, ignore_errors=True)
            except Exception:
                pass

        # 下载文件索引（与下载模块联动）
        if getattr(self, '_download_info', None):
            doc.add_heading('下载文件索引', level=1)
            for i, f in enumerate(self._download_info, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {f.get('filename', '')}").bold = True
                doc.add_paragraph(f"   URL: {f.get('url', '')}")
                doc.add_paragraph(f"   本地路径: {f.get('local_path', '')}")

        doc.save(str(path))
        return str(path)

    def _load_json(self, path: Path) -> List[Dict]:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            return data
        return [data]

    def _load_csv(self, path: Path) -> List[Dict]:
        with open(path, "r", encoding="utf-8") as f:
            reader = csv.DictReader(f)
            return list(reader)

    def _load_sqlite(self, path: Path) -> List[Dict]:
        conn = sqlite3.connect(str(path))
        conn.row_factory = sqlite3.Row
        cursor = conn.cursor()

        # 获取表名
        cursor.execute("SELECT name FROM sqlite_master WHERE type='table'")
        tables = cursor.fetchall()
        if not tables:
            conn.close()
            return []

        table_name = tables[0]["name"]
        cursor.execute(f"SELECT * FROM {table_name}")
        rows = cursor.fetchall()
        result = [dict(row) for row in rows]

        conn.close()
        return result

    def _load_markdown(self, path: Path) -> List[Dict]:
        """简单的 Markdown 表格解析"""
        content = path.read_text(encoding="utf-8")
        lines = content.strip().split("\n")

        if len(lines) < 3:
            return []

        # 解析表头
        header_line = lines[0].strip()
        if not header_line.startswith("|"):
            return []

        headers = [h.strip() for h in header_line.split("|")[1:-1]]

        # 跳过分隔行
        data_lines = lines[2:]
        result = []
        for line in data_lines:
            line = line.strip()
            if not line.startswith("|"):
                continue
            values = [v.strip().replace("\\|", "|") for v in line.split("|")[1:-1]]
            item = dict(zip(headers, values))
            result.append(item)

        return result

    def _load_txt(self, path: Path) -> List[Dict]:
        """从纯文本加载 - 简单解析"""
        content = path.read_text(encoding="utf-8")
        if not content.strip():
            return []

        result = []
        current_item = {}
        current_key = None
        current_value = []

        for line in content.split("\n"):
            line = line.rstrip()
            
            # 分隔线
            if line.startswith("=" * 10):
                if current_key and current_value:
                    current_item[current_key] = "\n".join(current_value)
                if current_item:
                    result.append(current_item)
                current_item = {}
                current_key = None
                current_value = []
                continue
            
            # 字段行
            if ":" in line and not line.startswith(" "):
                # 保存之前的字段
                if current_key and current_value:
                    current_item[current_key] = "\n".join(current_value)
                
                # 新字段
                key, _, value = line.partition(":")
                current_key = key.strip()
                current_value = [value.strip()] if value.strip() else []
            else:
                # 续行
                if current_key:
                    current_value.append(line)

        # 最后一条
        if current_key and current_value:
            current_item[current_key] = "\n".join(current_value)
        if current_item:
            result.append(current_item)

        return result
